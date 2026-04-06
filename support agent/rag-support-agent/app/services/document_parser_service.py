"""
Document parser service for handling multiple document types.

Supports:
- PDF (via pdfplumber)
- DOCX (via python-docx)
- Plain text

Each parser returns normalized text content that can be chunked uniformly.
"""

import logging
import os
from pathlib import Path
from typing import Tuple, Optional
from enum import Enum

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    UNKNOWN = "unknown"


class DocumentMetadata:
    """Metadata extracted from documents."""
    
    def __init__(
        self,
        filename: str,
        document_type: DocumentType,
        num_pages: Optional[int] = None,
        num_words: int = 0,
        source_path: Optional[str] = None,
    ):
        self.filename = filename
        self.document_type = document_type
        self.num_pages = num_pages
        self.num_words = num_words
        self.source_path = source_path
    
    def to_dict(self) -> dict:
        """Convert metadata to dictionary."""
        return {
            "filename": self.filename,
            "document_type": self.document_type.value,
            "num_pages": self.num_pages,
            "num_words": self.num_words,
            "source_path": self.source_path,
        }


class DocumentParser:
    """Base class for document parsers."""
    
    def parse(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """
        Parse document and return content + metadata.
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (content, metadata)
        """
        raise NotImplementedError


class PDFParser(DocumentParser):
    """PDF document parser using pdfplumber."""
    
    def parse(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """
        Parse PDF file and extract text.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (content, metadata)
        """
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber required for PDF parsing. Install: pip install pdfplumber")
        
        try:
            content_parts = []
            num_pages = 0
            
            with pdfplumber.open(file_path) as pdf:
                num_pages = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        text = page.extract_text()
                        if text:
                            # Add page marker for context
                            content_parts.append(f"[Page {page_num}]\n{text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue
            
            content = "\n\n".join(content_parts)
            num_words = len(content.split())
            
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                document_type=DocumentType.PDF,
                num_pages=num_pages,
                num_words=num_words,
                source_path=file_path,
            )
            
            logger.info(f"Parsed PDF: {metadata.filename} ({num_pages} pages, {num_words} words)")
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise


class DOCXParser(DocumentParser):
    """DOCX document parser using python-docx."""
    
    def parse(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """
        Parse DOCX file and extract text.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (content, metadata)
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for DOCX parsing. Install: pip install python-docx")
        
        try:
            doc = Document(file_path)
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    content_parts.append("\n".join(table_text))
            
            content = "\n\n".join(content_parts)
            num_words = len(content.split())
            
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                document_type=DocumentType.DOCX,
                num_pages=None,  # DOCX doesn't have pages
                num_words=num_words,
                source_path=file_path,
            )
            
            logger.info(f"Parsed DOCX: {metadata.filename} ({num_words} words)")
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise


class TextParser(DocumentParser):
    """Plain text file parser."""
    
    def parse(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """
        Parse plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Tuple of (content, metadata)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            num_words = len(content.split())
            
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                document_type=DocumentType.TXT,
                num_words=num_words,
                source_path=file_path,
            )
            
            logger.info(f"Parsed text file: {metadata.filename} ({num_words} words)")
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {e}")
            raise


class DocumentParserFactory:
    """Factory for creating appropriate parser based on file type."""
    
    _parsers = {
        DocumentType.PDF: PDFParser,
        DocumentType.DOCX: DOCXParser,
        DocumentType.TXT: TextParser,
    }
    
    @staticmethod
    def get_document_type(file_path: str) -> DocumentType:
        """
        Detect document type from file extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            DocumentType enum value
        """
        ext = Path(file_path).suffix.lower()
        
        type_map = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
        }
        
        return type_map.get(ext, DocumentType.UNKNOWN)
    
    @staticmethod
    def parse(file_path: str) -> Tuple[str, DocumentMetadata]:
        """
        Parse document with appropriate parser.
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (content, metadata)
            
        Raises:
            ValueError: If document type is unsupported
        """
        # Validate file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detect document type
        doc_type = DocumentParserFactory.get_document_type(file_path)
        
        if doc_type not in DocumentParserFactory._parsers:
            raise ValueError(f"Unsupported document type: {doc_type}. Supported: {list(DocumentParserFactory._parsers.keys())}")
        
        # Get appropriate parser and parse
        parser_class = DocumentParserFactory._parsers[doc_type]
        parser = parser_class()
        
        return parser.parse(file_path)
    
    @staticmethod
    def supports(file_path: str) -> bool:
        """Check if document type is supported."""
        doc_type = DocumentParserFactory.get_document_type(file_path)
        return doc_type in DocumentParserFactory._parsers
